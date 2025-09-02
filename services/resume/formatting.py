"""
DOCX and PDF formatting for optimized resumes.
Creates professional document outputs from structured resume data.
"""

import io
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import asyncio

from .schemas import OptimizedResume, ExperienceItem


class ModernResumeFormatter:
    """Creates modern, ATS-friendly DOCX resumes."""
    
    def __init__(self, template_style: str = "professional"):
        self.template_style = template_style
        self.doc = Document()
        self._setup_document_styles()
    
    def _setup_document_styles(self):
        """Set up custom styles for the document."""
        
        # Document margins (narrow for more content)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Create custom styles
        styles = self.doc.styles
        
        # Header style (name and contact)
        if 'ResumeHeader' not in [s.name for s in styles]:
            header_style = styles.add_style('ResumeHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Calibri'
            header_style.font.size = Pt(18)
            header_style.font.bold = True
            header_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Dark blue
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_style.paragraph_format.space_after = Pt(6)
        
        # Contact info style
        if 'ContactInfo' not in [s.name for s in styles]:
            contact_style = styles.add_style('ContactInfo', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.name = 'Calibri'
            contact_style.font.size = Pt(11)
            contact_style.font.color.rgb = RGBColor(0x64, 0x64, 0x64)  # Gray
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_style.paragraph_format.space_after = Pt(12)
        
        # Section header style
        if 'SectionHeader' not in [s.name for s in styles]:
            section_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Calibri'
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
            # Add bottom border
            self._add_bottom_border(section_style)
        
        # Job header style
        if 'JobHeader' not in [s.name for s in styles]:
            job_style = styles.add_style('JobHeader', WD_STYLE_TYPE.PARAGRAPH)
            job_style.font.name = 'Calibri'
            job_style.font.size = Pt(12)
            job_style.font.bold = True
            job_style.paragraph_format.space_before = Pt(8)
            job_style.paragraph_format.space_after = Pt(2)
        
        # Bullet point style
        if 'BulletPoint' not in [s.name for s in styles]:
            bullet_style = styles.add_style('BulletPoint', WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.font.name = 'Calibri'
            bullet_style.font.size = Pt(11)
            bullet_style.paragraph_format.left_indent = Inches(0.25)
            bullet_style.paragraph_format.space_after = Pt(3)
        
        # Skills style
        if 'SkillsList' not in [s.name for s in styles]:
            skills_style = styles.add_style('SkillsList', WD_STYLE_TYPE.PARAGRAPH)
            skills_style.font.name = 'Calibri'
            skills_style.font.size = Pt(11)
            skills_style.paragraph_format.space_after = Pt(3)
    
    def _add_bottom_border(self, style):
        """Add bottom border to paragraph style."""
        # This is a simplified border - python-docx border handling is complex
        # For MVP, we'll skip the border and use spacing instead
        style.paragraph_format.space_after = Pt(8)
    
    def format_resume(
        self,
        optimized_resume: OptimizedResume,
        contact_info: Dict[str, str] = None,
        template_options: Dict[str, Any] = None
    ) -> Document:
        """
        Format optimized resume into professional DOCX document.
        
        Args:
            optimized_resume: Structured resume data
            contact_info: Contact information (name, email, phone, etc.)
            template_options: Additional formatting options
            
        Returns:
            Formatted Document object
        """
        
        print(f"📄 Formatting resume to DOCX...")
        
        if not contact_info:
            contact_info = {}
        
        # 1. Header (Name and Contact)
        self._add_header(contact_info)
        
        # 2. Professional Summary
        if optimized_resume.summary:
            self._add_summary(optimized_resume.summary)
        
        # 3. Experience Section
        if optimized_resume.experience:
            self._add_experience_section(optimized_resume.experience)
        
        # 4. Skills Section
        if optimized_resume.skills or optimized_resume.skills_to_add:
            all_skills = (optimized_resume.skills or []) + (optimized_resume.skills_to_add or [])
            self._add_skills_section(list(dict.fromkeys(all_skills)))  # Remove duplicates
        
        # 5. Education Section (if provided)
        if optimized_resume.education:
            self._add_education_section(optimized_resume.education)
        
        print(f"✅ DOCX formatting complete")
        return self.doc
    
    def _add_header(self, contact_info: Dict[str, str]):
        """Add name and contact information header."""
        
        # Name
        name = contact_info.get('name', 'Your Name')
        name_para = self.doc.add_paragraph(name, style='ResumeHeader')
        
        # Contact info line
        contact_parts = []
        if contact_info.get('email'):
            contact_parts.append(contact_info['email'])
        if contact_info.get('phone'):
            contact_parts.append(contact_info['phone'])
        if contact_info.get('location'):
            contact_parts.append(contact_info['location'])
        if contact_info.get('linkedin'):
            contact_parts.append(f"LinkedIn: {contact_info['linkedin']}")
        
        if contact_parts:
            contact_line = ' | '.join(contact_parts)
            self.doc.add_paragraph(contact_line, style='ContactInfo')
    
    def _add_summary(self, summary_text: str):
        """Add professional summary section."""
        
        self.doc.add_paragraph('PROFESSIONAL SUMMARY', style='SectionHeader')
        summary_para = self.doc.add_paragraph(summary_text)
        summary_para.style.font.name = 'Calibri'
        summary_para.style.font.size = Pt(11)
        summary_para.style.paragraph_format.space_after = Pt(8)
    
    def _add_experience_section(self, experience_items: List[ExperienceItem]):
        """Add experience section with job entries."""
        
        self.doc.add_paragraph('EXPERIENCE', style='SectionHeader')
        
        for exp_item in experience_items:
            # Job header (Role | Company | Dates)
            job_header_parts = [exp_item.role]
            if exp_item.company:
                job_header_parts.append(exp_item.company)
            if exp_item.dates:
                job_header_parts.append(exp_item.dates)
            
            job_header = ' | '.join(job_header_parts)
            self.doc.add_paragraph(job_header, style='JobHeader')
            
            # Bullet points
            for bullet in exp_item.bullets:
                bullet_para = self.doc.add_paragraph(f"• {bullet}", style='BulletPoint')
    
    def _add_skills_section(self, skills: List[str]):
        """Add technical skills section."""
        
        self.doc.add_paragraph('TECHNICAL SKILLS', style='SectionHeader')
        
        # Group skills into categories if possible
        categorized_skills = self._categorize_skills(skills)
        
        if categorized_skills:
            # Display by category
            for category, skill_list in categorized_skills.items():
                if skill_list:
                    category_line = f"{category.title()}: {', '.join(skill_list)}"
                    self.doc.add_paragraph(category_line, style='SkillsList')
        else:
            # Simple comma-separated list
            skills_line = ', '.join(skills)
            self.doc.add_paragraph(skills_line, style='SkillsList')
    
    def _add_education_section(self, education_items: List[Dict[str, str]]):
        """Add education section."""
        
        self.doc.add_paragraph('EDUCATION', style='SectionHeader')
        
        for edu_item in education_items:
            edu_parts = []
            if edu_item.get('degree'):
                edu_parts.append(edu_item['degree'])
            if edu_item.get('school'):
                edu_parts.append(edu_item['school'])
            if edu_item.get('year') and edu_item['year']:
                # Apply age signal scrubbing
                try:
                    year = int(edu_item['year'])
                    current_year = 2025
                    if current_year - year <= 15:  # Only show recent graduations
                        edu_parts.append(str(year))
                except:
                    pass  # Skip invalid years
            
            if edu_parts:
                edu_line = ' | '.join(edu_parts)
                self.doc.add_paragraph(edu_line, style='SkillsList')
    
    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize skills for better organization."""
        
        from .keywords import TechnicalSkills
        
        categorized = {
            'Languages': [],
            'Frameworks': [],
            'Cloud & DevOps': [],
            'Databases': [],
            'Other': []
        }
        
        for skill in skills:
            skill_lower = skill.lower()
            category = TechnicalSkills.categorize_skill(skill)
            
            if category == 'languages':
                categorized['Languages'].append(skill)
            elif category in ['frameworks']:
                categorized['Frameworks'].append(skill)
            elif category in ['cloud', 'devops']:
                categorized['Cloud & DevOps'].append(skill)
            elif category == 'databases':
                categorized['Databases'].append(skill)
            else:
                categorized['Other'].append(skill)
        
        # Remove empty categories
        return {k: v for k, v in categorized.items() if v}
    
    def to_bytes(self) -> bytes:
        """Convert document to bytes for download/storage."""
        
        bio = io.BytesIO()
        self.doc.save(bio)
        bio.seek(0)
        return bio.getvalue()


class ResumeDocumentBuilder:
    """Main interface for building resume documents."""
    
    @staticmethod
    async def create_docx(
        optimized_resume: OptimizedResume,
        contact_info: Dict[str, str] = None,
        template_style: str = "professional"
    ) -> bytes:
        """
        Create DOCX document from optimized resume.
        
        Args:
            optimized_resume: Structured resume data
            contact_info: Contact information
            template_style: Template style to use
            
        Returns:
            DOCX file as bytes
        """
        
        formatter = ModernResumeFormatter(template_style)
        document = formatter.format_resume(optimized_resume, contact_info)
        
        return formatter.to_bytes()
    
    @staticmethod
    async def create_pdf(
        optimized_resume: OptimizedResume,
        contact_info: Dict[str, str] = None,
        template_style: str = "professional"
    ) -> bytes:
        """
        Create PDF document from optimized resume using WeasyPrint.
        
        Args:
            optimized_resume: Structured resume data
            contact_info: Contact information
            template_style: Template style to use
            
        Returns:
            PDF file as bytes
        """
        
        # First create HTML version
        html_content = HTMLResumeFormatter.format_to_html(
            optimized_resume, contact_info, template_style
        )
        
        # Convert to PDF using WeasyPrint
        from weasyprint import HTML, CSS
        
        # Custom CSS for PDF
        pdf_css = CSS(string="""
            @page {
                size: letter;
                margin: 0.75in;
            }
            body {
                font-family: 'Calibri', Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.4;
                color: #333;
            }
            .header {
                text-align: center;
                margin-bottom: 20px;
            }
            .name {
                font-size: 18pt;
                font-weight: bold;
                color: #2C3E50;
                margin-bottom: 5px;
            }
            .contact {
                font-size: 11pt;
                color: #666;
                margin-bottom: 15px;
            }
            .section-header {
                font-size: 14pt;
                font-weight: bold;
                color: #2C3E50;
                border-bottom: 1px solid #ccc;
                margin-top: 15px;
                margin-bottom: 8px;
                padding-bottom: 3px;
            }
            .job-header {
                font-weight: bold;
                margin-top: 10px;
                margin-bottom: 3px;
            }
            .bullet {
                margin-left: 20px;
                margin-bottom: 3px;
            }
            .skills {
                line-height: 1.5;
            }
        """)
        
        pdf_bytes = HTML(string=html_content).write_pdf(stylesheets=[pdf_css])
        
        return pdf_bytes


class HTMLResumeFormatter:
    """Creates HTML version of resume for PDF conversion."""
    
    @staticmethod
    def format_to_html(
        optimized_resume: OptimizedResume,
        contact_info: Dict[str, str] = None,
        template_style: str = "professional"
    ) -> str:
        """
        Format resume to HTML for PDF conversion.
        
        Args:
            optimized_resume: Structured resume data
            contact_info: Contact information
            template_style: Template style
            
        Returns:
            HTML string
        """
        
        if not contact_info:
            contact_info = {}
        
        html_parts = ['<!DOCTYPE html>', '<html>', '<head>']
        html_parts.append('<meta charset="utf-8">')
        html_parts.append('<title>Resume</title>')
        html_parts.extend(['</head>', '<body>'])
        # Header
        html_parts.append('<div class="header">')
        name = contact_info.get('name', 'Your Name')
        html_parts.append(f'<div class="name">{name}</div>')
        
        # Contact info
        contact_parts = []
        if contact_info.get('email'):
            contact_parts.append(contact_info['email'])
        if contact_info.get('phone'):
            contact_parts.append(contact_info['phone'])
        if contact_info.get('location'):
            contact_parts.append(contact_info['location'])
        
        if contact_parts:
            contact_line = ' | '.join(contact_parts)
            html_parts.append(f'<div class="contact">{contact_line}</div>')
        
        html_parts.append('</div>')
        
        # Professional Summary
        if optimized_resume.summary:
            html_parts.append('<div class="section-header">PROFESSIONAL SUMMARY</div>')
            html_parts.append(f'<p>{optimized_resume.summary}</p>')
        
        # Experience
        if optimized_resume.experience:
            html_parts.append('<div class="section-header">EXPERIENCE</div>')
            
            for exp_item in optimized_resume.experience:
                # Job header
                job_parts = [exp_item.role]
                if exp_item.company:
                    job_parts.append(exp_item.company)
                if exp_item.dates:
                    job_parts.append(exp_item.dates)
                
                job_header = ' | '.join(job_parts)
                html_parts.append(f'<div class="job-header">{job_header}</div>')
                
                # Bullets
                for bullet in exp_item.bullets:
                    html_parts.append(f'<div class="bullet">• {bullet}</div>')
        
        # Skills
        all_skills = (optimized_resume.skills or []) + (optimized_resume.skills_to_add or [])
        if all_skills:
            unique_skills = list(dict.fromkeys(all_skills))  # Remove duplicates
            html_parts.append('<div class="section-header">TECHNICAL SKILLS</div>')
            skills_line = ', '.join(unique_skills)
            html_parts.append(f'<div class="skills">{skills_line}</div>')
        
        # Education
        if optimized_resume.education:
            html_parts.append('<div class="section-header">EDUCATION</div>')
            for edu_item in optimized_resume.education:
                edu_parts = []
                if edu_item.get('degree'):
                    edu_parts.append(edu_item['degree'])
                if edu_item.get('school'):
                    edu_parts.append(edu_item['school'])
                if edu_item.get('year'):
                    edu_parts.append(edu_item['year'])
                
                if edu_parts:
                    edu_line = ' | '.join(edu_parts)
                    html_parts.append(f'<p>{edu_line}</p>')
        
        html_parts.extend(['</body>', '</html>'])
        
        return '\n'.join(html_parts)


# Sync wrappers
def create_docx_sync(
    optimized_resume: OptimizedResume,
    contact_info: Dict[str, str] = None,
    template_style: str = "professional"
) -> bytes:
    """Synchronous wrapper for create_docx."""
    return asyncio.run(ResumeDocumentBuilder.create_docx(optimized_resume, contact_info, template_style))


def create_pdf_sync(
    optimized_resume: OptimizedResume,
    contact_info: Dict[str, str] = None,
    template_style: str = "professional"
) -> bytes:
    """Synchronous wrapper for create_pdf."""
    return asyncio.run(ResumeDocumentBuilder.create_pdf(optimized_resume, contact_info, template_style))


class DocumentMetadata:
    """Metadata and file handling for generated documents."""
    
    @staticmethod
    def generate_filename(
        contact_name: str,
        job_title: str = None,
        file_format: str = "docx"
    ) -> str:
        """
        Generate appropriate filename for resume document.
        
        Args:
            contact_name: Person's name
            job_title: Job title (optional)
            file_format: File format (docx or pdf)
            
        Returns:
            Sanitized filename
        """
        
        # Sanitize name
        safe_name = re.sub(r'[^\w\s-]', '', contact_name).strip()
        safe_name = re.sub(r'\s+', '_', safe_name)
        
        # Add job title if provided
        if job_title:
            safe_job = re.sub(r'[^\w\s-]', '', job_title).strip()
            safe_job = re.sub(r'\s+', '_', safe_job)
            filename = f"{safe_name}_Resume_{safe_job}.{file_format}"
        else:
            filename = f"{safe_name}_Resume_Optimized.{file_format}"
        
        return filename
    
    @staticmethod
    def get_file_info(document_bytes: bytes, file_format: str) -> Dict[str, Any]:
        """Get metadata about generated document file."""
        
        return {
            'size_bytes': len(document_bytes),
            'size_kb': round(len(document_bytes) / 1024, 2),
            'format': file_format,
            'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_format == 'docx' else 'application/pdf',
            'generated_at': datetime.now().isoformat()
        }